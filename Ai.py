import requests
from docx import Document

def read_file(file_path):
    """Reads the content of a text file."""
    try:
        with open(file_path, 'r') as file:
            content = file.read()
        return content
    except FileNotFoundError:
        print("File not found. Please check the file path.")
        return None

def summarize_text(text):
    """Summarizes the text using a hypothetical Gemini API."""
    # Assuming there's a Gemini API endpoint that provides text summarization.
    api_url = "https://api.gemini.example/summarize"
    headers = {
        "Authorization": "Bearer AIzaSyCrmksb7T-Zm1PdPJUQwFJsVmJYQecb3CU",
        "Content-Type": "application/json"
    }
    payload = {
        "text": text
    }
    
    try:
        response = requests.post(api_url, json=payload, headers=headers)
        response.raise_for_status()
        summary = response.json().get("summary", "")
        return summary
    except requests.exceptions.RequestException as e:
        print(f"API request failed: {e}")
        return None

def create_docx(summary, output_file):
    """Creates a .docx file with the summarized text."""
    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary)
    doc.save(output_file)
    print(f"Summary saved to {output_file}")

def main():
    file_path = './New_Heuristic_Algorithm_to_improve_the_Minimax_for_Gomoku_Artificial_Intelligence (1).pdf'  # Specify the path to your text file
    output_file = 'summary.docx'  # Output .docx file
    
    # Step 1: Read the file
    text = read_file(file_path)
    if text is None:
        return
    
    # Step 2: Summarize the text
    summary = summarize_text(text)
    if summary is None:
        return
    
    # Step 3: Create a .docx file with the summary
    create_docx(summary, output_file)

if __name__ == "__main__":
    main()
